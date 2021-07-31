import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

doc=docx.Document()
#person_information
name=input('what is your name --> ')
email=input('what is your Email --> ')
add=input('what is your Address --> ')
phone=input('what is your Phone --> ')
college=input('Your College Name --> ')
#College Location and Grad
Location=input('Location of College --> ')
Location=Location+'\r'
y=input('enrollment year --> ')
col=college+'('+y+')'+'\r'
course=input('What Course --> ')
#any skill till now

Skill1=input('Skill1 --> ')
Skill1=Skill1+'\r'
desc=input('Description --> ')
#any job experience or any volunteering
Experience=input('Experience --> ')
Experience=Experience+'\r'
desc1=input('Description --> ')
#any course or certification
course1=input('ENROLLED COURSE  --> ')
course1=course1+'\r'
desc2=input('Description --> ')

#Social Contact
github=input('Github link --> ')
github=github+'\r'
facebook=input('Facebook link -->')
facebook=facebook+'\r'

#adding Personal Info to Resume
name=name+'\r'
stra=email+'\r'+add+'\r'+phone
a=doc.add_paragraph()
a.add_run(name).bold=True
a.add_run(stra)
a.alignment=WD_ALIGN_PARAGRAPH.RIGHT
c=doc.add_paragraph()

#adding your education qualifiacation
c.add_run('EDUCATION \r').bold=True
x=c.add_run('______________________________________________________________________________________________________\r')
font=x.font
font.color.rgb=RGBColor(0x42, 0x24, 0xE9)

c.add_run(col).bold=True
c.add_run(Location)
c.add_run(course)
#Your skills
b=doc.add_paragraph()
b.add_run('SKILLS \r').bold=True
x=b.add_run('______________________________________________________________________________________________________\r')
font=x.font
font.color.rgb=RGBColor(0x42, 0x24, 0xE9)
b.add_run(Skill1).bold=True
b.add_run(desc)

d=doc.add_paragraph()
d.add_run('EXPERIENCE \r').bold=True
x=d.add_run('______________________________________________________________________________________________________\r')
font=x.font
font.color.rgb=RGBColor(0x42, 0x24, 0xE9)
d.add_run(Experience).bold=True
d.add_run(desc1)
#Your Courses
e=doc.add_paragraph()
e.add_run('ENROLLED COURSES \r').bold=True
x=e.add_run('______________________________________________________________________________________________________\r')
font=x.font
font.color.rgb=RGBColor(0x42, 0x24, 0xE9)
e.add_run(course1).bold=True
e.add_run(desc2)
#Contact Details
f=doc.add_paragraph()
f.add_run('CONTACT DETAILS \r').bold=True
x=f.add_run('______________________________________________________________________________________________________\r')
font=x.font
font.color.rgb=RGBColor(0x42, 0x24, 0xE9)
f.add_run('Facebook:').bold=True
f.add_run(facebook)
f.add_run('Github:').bold=True
f.add_run(github)


doc.save('C:\\Users\\Kumar\\Desktop\\Resume.docx') #you can add your Directory 

